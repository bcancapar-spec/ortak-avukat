# -*- coding: utf-8 -*-
"""oa-dilekce / udf_yaz.py için ALTIN VAKA testleri.

Script'i dosya-yolundan (importlib.util) yükler — skill dizinleri paket değildir.

GÖREV D (v0.5.5 saha bulgusu B5, KRİTİK): eski hand-rolled zip/content.xml
motoru (`udf_uret`/`udf_yaz`/`udf_metni_geri_oku`/`--yerel-motor`) UYAP'ta
AÇILMAYAN bir `.udf` üretiyordu — bu sınıf TAMAMEN KALDIRILDI. Script artık
TEK yazma hattına sahiptir: md → UDF-HTML (md_udf_html.py) → gerçek
`npx -y udf-cli@latest html2udf`. Bu yüzden bu dosyadaki testler de:
  - `udf_dogrula()` (mekanik GEÇERLİLİK KAPISI) testleri artık kendi sentetik
    ZIP/content.xml'lerini DOĞRUDAN `zipfile` ile kurar (üretici script'e
    bağımlı DEĞİLDİR — bu bir test fixture'ıdır, "elle UDF üretimi" ÖZELLİĞİ
    DEĞİLDİR).
  - Gerçek `npx udf-cli html2udf` uçtan-uca testleri `npx_kullanilabilir_mi()`
    ile KULLANILABİLİRSE koşar, DEĞİLSE nazikçe `pytest.skip` eder (hiçbir
    makinede KIRILMAZ).
  - FAIL-CLOSED testleri (npx yok/başarısız → dosya YAZILMAZ, exit != 0, net
    hata mesajı) ağsız/deterministiktir, her zaman koşar.
"""
import importlib.util
import pathlib
import shutil
import subprocess
import sys
import zipfile

import pytest

REPO = pathlib.Path(__file__).resolve().parents[1]
SCRIPT = REPO / "plugins" / "ortak-avukat" / "skills" / "oa-dilekce" / "scripts" / "udf_yaz.py"
MD_HTML_SCRIPT = REPO / "plugins" / "ortak-avukat" / "skills" / "oa-dilekce" / "scripts" / "md_udf_html.py"


def _load():
    assert SCRIPT.is_file(), f"udf_yaz.py bulunamadı: {SCRIPT}"
    spec = importlib.util.spec_from_file_location("udf_yaz", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


uy = _load()


# ── ALTIN KURAL: elle-zip motoru KALDIRILDI (regresyon kilidi) ──────────────

def test_yerel_motor_fonksiyonlari_artik_yok():
    """B5 düzeltmesi: hand-rolled zip/content.xml üreticisi (`udf_uret`,
    `udf_yaz` yazıcı fonksiyonu, `udf_metni_geri_oku`, `cdata_guvenli`,
    `md_satir_duzlestir`) script'ten TAMAMEN KALDIRILMIŞ olmalı — bunlardan
    biri geri gelirse "sessizce eski yola düşme" riski de geri gelir."""
    for ad in ("udf_uret", "udf_metni_geri_oku", "cdata_guvenli",
               "md_satir_duzlestir"):
        assert not hasattr(uy, ad), (
            f"'{ad}' hâlâ mevcut — elle-zip motoru tam kaldırılmamış olabilir")


def test_yerel_motor_cli_bayragi_artik_yok():
    """--yerel-motor / --format-id argparse'tan kaldırılmış olmalı."""
    cp = subprocess.run([sys.executable, str(SCRIPT), "--help"],
                         capture_output=True, text=True, encoding="utf-8", errors="replace")
    assert "--yerel-motor" not in cp.stdout
    assert "--format-id" not in cp.stdout


def test_md2udf_kodda_hic_gecmiyor():
    """ALTIN VAKA (rehber kuralı): 'md2udf' hiçbir zaman bir komut/argüman
    olarak çağrılmamalı — yalnız dokümantasyon amaçlı ('ASLA md2udf kullanma'
    uyarısı) serbest metinde geçebilir; kod, bunu quoted bir CLI argümanı
    olarak ASLA üretmemeli. html2udf ise gerçek yazıcı olarak MUTLAKA geçmeli."""
    kaynak_udf_yaz = SCRIPT.read_text(encoding="utf-8")
    kaynak_md_html = MD_HTML_SCRIPT.read_text(encoding="utf-8")
    for kaynak, ad in ((kaynak_udf_yaz, "udf_yaz.py"), (kaynak_md_html, "md_udf_html.py")):
        assert '"md2udf"' not in kaynak and "'md2udf'" not in kaynak, (
            f"{ad} içinde 'md2udf' TIRNAKLI bir literal (muhtemel CLI argümanı) bulundu")
    assert '"html2udf"' in kaynak_udf_yaz or "'html2udf'" in kaynak_udf_yaz


# ── udf_dogrula(): UDF GEÇERLİLİK KAPISI (mekanik, hüküm YOK) ───────────────
#
# NOT (v0.5.5.2): bu bölümdeki testler kapının YAPISAL bacağını (zip/XML/CDATA/
# offset döşemesi) sınar ve bu yüzden `resmi_okuyucu=False` ile çağırır — 5.
# bacak (`udf-cli udf2md` ile dışarıdan tanık) ağ+oturum ister ve KENDİ testleri
# `tests/test_udf_resmi_okuyucu.py`de sahte okuyucu enjekte edilerek koşar.
# Burada açık bırakmak, yapısal regresyon testlerini ağ durumuna bağımlı
# kılardı (çevrimdışı koşuda yanlış kırmızı).
# NOT: aşağıdaki yardımcı, script'in ürettiği bir UDF'i DEĞİL, testin kendi
# kurduğu sentetik bir ZIP/content.xml'i kullanır — script artık UDF üretmiyor
# (yalnız npx'e devrediyor), ama udf_dogrula() üretilen HERHANGİ bir UDF
# üzerinde (gerçek udf-cli çıktısı dâhil) bağımsız çalışabilmelidir.

def _sentetik_udf_yaz(tmp_path, metin="Birinci satır.\nİkinci satır.\n", ad="gecerli.udf"):
    satirlar = metin.split("\n")
    if satirlar and satirlar[-1] == "":
        satirlar = satirlar[:-1]
    parcalar = [s + "\n" for s in satirlar]
    tam = "".join(parcalar)

    xml = ['<?xml version="1.0" encoding="UTF-8"?>',
           '<template format_id="1.8">',
           '<content><![CDATA[' + tam.replace("]]>", "]]]]><![CDATA[>") + ']]></content>',
           '<properties><pageFormat mediaSizeName="1" leftMargin="70.866" '
           'rightMargin="70.866" topMargin="70.866" bottomMargin="70.866" '
           'paperOrientation="1"/></properties>',
           '<elements resolver="hvl-default">']
    imlec = 0
    for parca in parcalar:
        u16 = uy.utf16_uzunluk(parca)
        xml.append('<paragraph Alignment="3"><content startOffset="%d" length="%d"/></paragraph>'
                    % (imlec, u16))
        imlec += u16
    xml.append('</elements>')
    xml.append('<styles><style name="default" family="Times New Roman" size="12"/></styles>')
    xml.append('</template>')
    xml_str = "\n".join(xml) + "\n"

    cikti = tmp_path / ad
    with zipfile.ZipFile(str(cikti), "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("content.xml", xml_str.encode("utf-8"))
    return cikti


def test_udf_dogrula_gecerli_dosyada_GECERLI_doner(tmp_path):
    """Doğru şekilde kurulmuş bir UDF (test fixture'ı — script'in kendisi
    değil) `udf_dogrula` ile GEÇERLİ dönmeli."""
    cikti = _sentetik_udf_yaz(tmp_path, "# Dava Dilekçesi\n\nSayın Mahkeme,\n\nArz ederiz.\n")
    sonuc = uy.udf_dogrula(str(cikti), resmi_okuyucu=False)
    assert sonuc["gecerli"] is True
    assert sonuc["hatalar"] == []
    assert sonuc["content_xml_var"] is True
    assert sonuc["xml_iyi_bicimli"] is True
    assert sonuc["cdata_bulundu"] is True
    assert sonuc["offsetler_tutarli"] is True
    assert sonuc["paragraf_sayisi"] > 0


def test_udf_dogrula_bozuk_zip_yakalar(tmp_path):
    """ZIP olmayan / bozuk bir dosya GEÇERSİZ dönmeli, exception fırlatmamalı."""
    sahte = tmp_path / "bozuk.udf"
    sahte.write_bytes(b"bu bir zip arsivi degil")
    sonuc = uy.udf_dogrula(str(sahte), resmi_okuyucu=False)
    assert sonuc["gecerli"] is False
    assert sonuc["hatalar"]
    assert sonuc["content_xml_var"] is False


def test_udf_dogrula_olmayan_dosya_yakalar(tmp_path):
    """Hiç var olmayan bir yol GEÇERSİZ dönmeli (FileNotFoundError yutulur)."""
    sonuc = uy.udf_dogrula(str(tmp_path / "yok.udf"), resmi_okuyucu=False)
    assert sonuc["gecerli"] is False
    assert sonuc["hatalar"]


def test_udf_dogrula_content_xml_eksik_zip_yakalar(tmp_path):
    """Geçerli bir ZIP ama içinde content.xml yoksa GEÇERSİZ olmalı."""
    sahte = tmp_path / "icersiz.udf"
    with zipfile.ZipFile(str(sahte), "w") as z:
        z.writestr("baska.txt", "ilgisiz içerik")
    sonuc = uy.udf_dogrula(str(sahte), resmi_okuyucu=False)
    assert sonuc["gecerli"] is False
    assert sonuc["content_xml_var"] is False
    assert any("content.xml" in h for h in sonuc["hatalar"])


def test_udf_dogrula_bozuk_xml_yakalar(tmp_path):
    """content.xml var ama iyi biçimli XML değilse GEÇERSİZ olmalı."""
    sahte = tmp_path / "bozukxml.udf"
    with zipfile.ZipFile(str(sahte), "w") as z:
        z.writestr("content.xml", "<template><content><![CDATA[eksik kapanis")
    sonuc = uy.udf_dogrula(str(sahte), resmi_okuyucu=False)
    assert sonuc["gecerli"] is False
    assert sonuc["content_xml_var"] is True
    assert sonuc["xml_iyi_bicimli"] is False


def test_udf_dogrula_tahrif_edilmis_offset_yakalar(tmp_path):
    """ALTIN VAKA: content.xml iyi biçimli ve CDATA doğru ama paragraf offset'i
    elle bozulmuşsa (ör. dosya sonradan tahrif edilmişse) offsetler_tutarli
    False dönmeli ve genel sonuç GEÇERSİZ olmalı — bu, yazımdan SONRA da
    tutarlılığı yakalayan bağımsız denetimdir."""
    cikti = _sentetik_udf_yaz(tmp_path, "Birinci satır.\nİkinci satır.\n")
    zf = zipfile.ZipFile(str(cikti))
    xml_ham = zf.read("content.xml").decode("utf-8")
    zf.close()
    # ilk paragrafın startOffset'ini bilerek bozuyoruz (0 → 5)
    bozuk_xml = xml_ham.replace('startOffset="0"', 'startOffset="5"', 1)
    assert bozuk_xml != xml_ham, "test kurulumu: değiştirilecek startOffset=\"0\" bulunamadı"
    with zipfile.ZipFile(str(cikti), "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("content.xml", bozuk_xml.encode("utf-8"))
    sonuc = uy.udf_dogrula(str(cikti), resmi_okuyucu=False)
    assert sonuc["gecerli"] is False
    assert sonuc["offsetler_tutarli"] is False
    assert sonuc["hatalar"]


# ── udf_dogrula(): tablo hücreleri yanlış GEÇERSİZ sayılmamalı ──────────────

def test_udf_dogrula_ic_ice_tablo_hucreli_belgeyi_GECERLI_sayar():
    """ALTIN VAKA: gerçek `udf-cli html2udf` çıktısında tablo hücreleri
    (<table><row><cell><paragraph><content .../>) üst-seviye <paragraph>'lardan
    AYRI bir dalda yaşar. Yalnız DİREKT paragraph/content arayan eski XPath
    bunları KAÇIRIR ve tamamen geçerli, tablolu bir belgeyi yanlışlıkla
    GEÇERSİZ işaretler — bu test o regresyonu kilitler (gerçek saha
    örneğinden alınmış sadeleştirilmiş content.xml)."""
    xml_ham = (
        '<?xml version="1.0" encoding="UTF-8" ?>\n'
        '<template format_id="1.8">\n'
        '<content><![CDATA[Baslik\nHucre1\nHucre2\n]]></content>'
        '<properties><pageFormat mediaSizeName="1" /></properties>\n'
        '<elements resolver="hvl-default">\n'
        '<paragraph Alignment="1"><content startOffset="0" length="7"/></paragraph>\n'
        '<table tableName="Sabit" columnCount="2">'
        '<row><cell><paragraph><content startOffset="7" length="7"/></paragraph></cell>'
        '<cell><paragraph><content startOffset="14" length="7"/></paragraph></cell></row>'
        '</table>\n'
        '</elements>\n'
        '<styles><style name="default" family="Times New Roman" size="12" /></styles>\n'
        '</template>\n'
    )
    import tempfile
    tmp = tempfile.mktemp(suffix=".udf")
    with zipfile.ZipFile(tmp, "w") as z:
        z.writestr("content.xml", xml_ham.encode("utf-8"))
    sonuc = uy.udf_dogrula(tmp, resmi_okuyucu=False)
    assert sonuc["paragraf_sayisi"] == 3, "üst paragraf + 2 tablo hücresi = 3 content elemanı"
    assert sonuc["offsetler_tutarli"] is True
    assert sonuc["gecerli"] is True, sonuc["hatalar"]


# ── _kok_coz(): --kok yol çözme yardımcısı ──────────────────────────────────

def test_kok_coz_goreli_yolu_koke_gore_cozer(tmp_path):
    sonuc = uy._kok_coz("alt/dosya.md", str(tmp_path))
    assert sonuc == str((tmp_path / "alt" / "dosya.md").resolve())


def test_kok_coz_mutlak_yolu_degistirmez(tmp_path):
    mutlak = str(tmp_path / "x.md")
    assert uy._kok_coz(mutlak, str(tmp_path / "baska")) == mutlak


def test_kok_coz_none_yol_none_doner():
    assert uy._kok_coz(None, "herhangi") is None


# ── md_html_uret(): kardeş script (md_udf_html.py) delegasyonu ─────────────

def test_md_html_uret_basligi_html_e_cevirir():
    html = uy.md_html_uret("# Başlık\n\nGövde metni.\n")
    assert "<strong>Başlık</strong>" in html
    assert "Gövde metni." in html


def test_md_html_uret_ham_modu_iletir():
    html = uy.md_html_uret("**kalın**\n", ham=True)
    assert "<strong>" not in html
    assert "**kalın**" in html


# ── UDF-uyumlu HTML şeması: pt birimi, kaçışsız <tab/>/<page-break/>, <br> ile
#    paragraf ayrımı YOK (rehber A.3/A.4 — bkz. uyap-belge-formatlari.md §3.1) ─

def test_html_uretiminde_br_ile_paragraf_ayrilmiyor():
    """Rehber kuralı: '<br><br>' ile paragraf ayırmak YASAK; her paragraf
    ayrı bir <p> bloğu olmalı."""
    html = uy.md_html_uret("Birinci paragraf.\n\nİkinci paragraf.\n")
    assert html.count("<p") >= 2
    assert "<br><br>" not in html


def test_html_uretiminde_pt_disinda_uzunluk_birimi_yok():
    """Üretilen inline-CSS'te px/em/rem/cm/mm/in gibi UDF'in güvenmediği
    birimler KULLANILMAMALI — yalnız pt (veya birimsiz sayı, html2udf'in
    varsayılanı) olmalı."""
    html = uy.md_html_uret("# Başlık\n\nGövde metni.\n")
    for yasak_birim in ("px", "em;", "rem;", "cm;", "mm;", "in;"):
        assert yasak_birim not in html, f"yasak birim '{yasak_birim}' HTML çıktısında bulundu"


def test_html_uretiminde_kullanici_metnindeki_ozel_karakterler_kacislanir():
    """md_udf_html.py şu an <tab/>/<page-break/> HİÇ ÜRETMİYOR (rehberin bu
    kontrolleri opsiyoneldir, script bunları emmiyor) — bu yüzden avukat
    taslağına düz metin olarak '<tab/>' yazarsa bu, işlevsel bir UDF tab-stop
    kontrolüne DÖNÜŞMEMELİ; HTML'de düz metin olarak GÖRÜNMESİ için ESCAPE
    edilmelidir (rehberin escape kuralı yalnız script'in KENDİSİ bir tab-stop/
    page-break ÜRETTİĞİNDE, o üretimi kaçışlamaması anlamına gelir — kullanıcı
    girdisindeki '<' / '>' genel HTML güvenliği için hep escape edilir)."""
    html = uy.md_html_uret("Kalem<tab/>Değer\n", ham=True)
    assert "&lt;tab/&gt;" in html
    assert "<tab/>" not in html


# ── npx yardımcıları: DETERMİNİST hata yolları (ağ gerekmez) ────────────────

def test_npx_kullanilabilir_mi_olmayan_komutla_false_doner():
    uygun, mesaj = uy.npx_kullanilabilir_mi(npx_yolu="oa-hic-boyle-bir-komut-yok-xyz")
    assert uygun is False
    assert mesaj


def test_npx_ile_udf_uret_olmayan_komutla_temiz_hata_doner(tmp_path):
    """npx bulunamadığında exception SIZDIRMAMALI — yapılandırılmış hata
    dönmeli VE hiçbir .udf dosyası YAZILMAMALI (FAIL-CLOSED — eski elle-zip
    yoluna SESSİZCE düşülmez)."""
    sonuc = uy.npx_ile_udf_uret(str(tmp_path / "x.html"), str(tmp_path / "x.udf"),
                                  npx_yolu="oa-hic-boyle-bir-komut-yok-xyz")
    assert sonuc["basarili"] is False
    assert sonuc["hata"]
    assert "FAIL-CLOSED" in sonuc["hata"]
    assert not (tmp_path / "x.udf").exists()


# ── CLI FAIL-CLOSED uçtan uca (tamamen ağsız/deterministik) ─────────────────

def test_cli_npx_yokken_udf_uretilmez_exit_farkli_sifir(tmp_path):
    """ALTIN VAKA (B5 regresyon kilidi): npx bulunamayan bir ortamda script
    --cikti dosyasını ASLA yazmamalı, exit kodu != 0 olmalı, ve stderr'de NET
    bir talimat (login / issue_cli_login_code) bulunmalı."""
    girdi = tmp_path / "taslak.md"
    girdi.write_text("# Dava Dilekçesi\n\nSayın Mahkeme, arz ederiz.\n", encoding="utf-8")
    cikti = tmp_path / "dilekce.udf"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--girdi", str(girdi), "--cikti", str(cikti),
         "--npx", "oa-hic-boyle-bir-komut-yok-xyz"],
        capture_output=True, text=True, encoding="utf-8", errors="replace")
    assert cp.returncode != 0
    assert not cikti.exists(), "npx yokken .udf dosyası YAZILMAMALIYDI (fail-closed ihlali)"
    assert "FAIL-CLOSED" in cp.stderr
    assert "login" in cp.stderr.lower()


def test_cli_npx_yokken_pdf_yine_de_denenir(tmp_path):
    """PDF önizlemesi UDF motorundan BAĞIMSIZDIR — npx/ağ olmasa bile
    (best-effort) üretilmeye çalışılmalı; UDF FAIL-CLOSED olsa da PDF
    bacağı ayrı bir artefakttır ve avukata en azından önizleme sağlar."""
    pytest.importorskip("fitz", reason="PyMuPDF kurulu değil")
    girdi = tmp_path / "taslak.md"
    girdi.write_text("# Dilekçe\n\nBirinci paragraf.\n", encoding="utf-8")
    cikti = tmp_path / "dilekce.udf"
    pdf = tmp_path / "dilekce.pdf"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--girdi", str(girdi), "--cikti", str(cikti),
         "--pdf", str(pdf), "--npx", "oa-hic-boyle-bir-komut-yok-xyz"],
        capture_output=True, text=True, encoding="utf-8", errors="replace")
    assert cp.returncode != 0
    assert not cikti.exists()
    assert pdf.is_file() and pdf.stat().st_size > 0


def test_cli_dogrula_modu_yazmadan_denetler(tmp_path):
    cikti = _sentetik_udf_yaz(tmp_path, "Metin.\n")
    cp = subprocess.run([sys.executable, str(SCRIPT), "--dogrula", str(cikti)],
                         capture_output=True, text=True, encoding="utf-8", errors="replace")
    assert cp.returncode == 0, cp.stdout + cp.stderr
    assert "GEÇERLİ" in cp.stdout


# ── BONUS YOL: docx2udf (rehber §5, elde hazır .docx/.pdf → UDF) ────────────

def test_docx2udf_cikis_kodu_tablosu_rehbere_uyuyor():
    assert uy.DOCX2UDF_CIKIS_KODU[0] == "Başarı"
    assert uy.DOCX2UDF_CIKIS_KODU[2].startswith("Giriş gerekli")
    assert uy.DOCX2UDF_CIKIS_KODU[3] == "Hesap yasaklı"
    assert uy.DOCX2UDF_CIKIS_KODU[5] == "Aylık kota bitti"


def test_docx2udf_ile_uret_npx_yokken_basarisiz_ve_dosya_yazilmaz(tmp_path):
    sahte_docx = tmp_path / "taslak.docx"
    sahte_docx.write_bytes(b"sahte docx icerigi")
    cikti = tmp_path / "cikti.udf"
    sonuc = uy.docx2udf_ile_uret(str(sahte_docx), str(cikti),
                                  npx_yolu="oa-hic-boyle-bir-komut-yok-xyz")
    assert sonuc["basarili"] is False
    assert sonuc["exit_kod"] is None
    assert not cikti.exists()


def test_cli_kaynak_docx_npx_yokken_exit_farkli_sifir(tmp_path):
    sahte_docx = tmp_path / "taslak.docx"
    sahte_docx.write_bytes(b"sahte docx icerigi")
    cikti = tmp_path / "cikti.udf"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--kaynak-docx", str(sahte_docx),
         "--cikti", str(cikti), "--npx", "oa-hic-boyle-bir-komut-yok-xyz"],
        capture_output=True, text=True, encoding="utf-8", errors="replace")
    assert cp.returncode != 0
    assert not cikti.exists()


def _docx2udf_hazir_mi():
    try:
        yol = shutil.which("npx")
        if yol is None:
            return False
        p = subprocess.run([yol, "-y", "docx2udf@latest", "whoami"],
                            capture_output=True, text=True, timeout=20)
        return p.returncode == 0
    except Exception:
        return False


@pytest.mark.skipif(not _docx2udf_hazir_mi(), reason="npx/docx2udf oturumu bu makinede kullanılamıyor")
def test_cli_kaynak_docx_gercek_donusum_ile_udf_uretir(tmp_path):
    pytest.importorskip("docx", reason="python-docx kurulu değil (test fixture için)")
    from docx import Document
    girdi = tmp_path / "taslak.docx"
    d = Document()
    d.add_heading("Dava Dilekcesi", level=1)
    d.add_paragraph("Sayin Mahkeme, arz ederiz.")
    d.save(str(girdi))

    cikti = tmp_path / "dilekce.udf"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--kaynak-docx", str(girdi), "--cikti", str(cikti)],
        capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=180)
    assert cp.returncode == 0, cp.stdout + cp.stderr
    assert cikti.is_file()
    zf = zipfile.ZipFile(str(cikti))
    assert "content.xml" in zf.namelist()


# ── varsayılan (npx udf-cli html2udf) hat — GERÇEK ağ/oturum gerektirir ─────
# `npx_kullanilabilir_mi()` YOKSA nazikçe atlanır (DOKUNULMAZLAR: testler
# hiçbir makinede KIRILMAZ); VARSA gerçek `udf-cli` çıktısını rehbere göre
# mekanik olarak doğrular.

def _npx_hazir_mi():
    try:
        uygun, _ = uy.npx_kullanilabilir_mi()
        return uygun
    except Exception:
        return False


@pytest.mark.skipif(not _npx_hazir_mi(), reason="npx/udf-cli oturumu bu makinede kullanılamıyor")
def test_cli_varsayilan_motor_sentetik_md_den_udf_uretir_ve_rehbere_gore_dogrular(tmp_path):
    """Kabul ölçütü: sentetik md → udf üret → zipfile ile aç, content.xml/
    properties şemasını rehbere göre doğrula."""
    girdi = tmp_path / "taslak.md"
    girdi.write_text(
        "# Dava Dilekçesi\n\nSayın Mahkeme,\n\n"
        "Müvekkilimiz **Ahmet Yılmaz** adına arz ederiz.\n\n"
        "| Kalem | Tutar |\n| --- | --- |\n| Vekalet | 1.000 TL |\n",
        encoding="utf-8")
    cikti = tmp_path / "dilekce.udf"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--girdi", str(girdi), "--cikti", str(cikti)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=180)
    assert cp.returncode == 0, cp.stdout + cp.stderr
    assert "GEÇERLİ" in cp.stdout

    # zipfile ile aç — rehberin "UDF bir ZIP arşividir" temel varsayımı
    zf = zipfile.ZipFile(str(cikti))
    assert "content.xml" in zf.namelist()
    xml_ham = zf.read("content.xml").decode("utf-8")

    # rehberin şemaya dair sabit beklentileri: template/format_id, CDATA
    # metin, properties/pageFormat, elements/paragraph/content, styles.
    assert "<template" in xml_ham and 'format_id="' in xml_ham
    assert "<![CDATA[" in xml_ham and "Ahmet Yılmaz" in xml_ham
    assert "<properties>" in xml_ham and "<pageFormat" in xml_ham
    assert "<elements" in xml_ham and "<paragraph" in xml_ham and "startOffset=" in xml_ham
    assert "<styles>" in xml_ham

    # mekanik geçerlilik kapısı da GEÇERLİ dönmeli (offset/CDATA tutarlılığı)
    sonuc = uy.udf_dogrula(str(cikti), resmi_okuyucu=False)
    assert sonuc["gecerli"] is True, sonuc["hatalar"]
    assert sonuc["paragraf_sayisi"] > 0


@pytest.mark.skipif(not _npx_hazir_mi(), reason="npx/udf-cli oturumu bu makinede kullanılamıyor")
def test_cli_varsayilan_motor_pdf_ile_birlikte_calisir(tmp_path):
    pytest.importorskip("fitz", reason="PyMuPDF kurulu değil")
    girdi = tmp_path / "taslak.md"
    girdi.write_text("# Dilekçe\n\nMetin.\n", encoding="utf-8")
    cikti = tmp_path / "dilekce.udf"
    pdf = tmp_path / "dilekce.pdf"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--girdi", str(girdi), "--cikti", str(cikti),
         "--pdf", str(pdf)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=180)
    assert cp.returncode == 0, cp.stdout + cp.stderr
    assert pdf.is_file() and pdf.stat().st_size > 0


@pytest.mark.skipif(not _npx_hazir_mi(), reason="npx/udf-cli oturumu bu makinede kullanılamıyor")
def test_cli_gecici_html_cikti_klasorune_dosya_birakmaz(tmp_path):
    """ALTIN VAKA (regresyon kilidi): --html AÇIKÇA verilmediğinde ara
    UDF-HTML SİSTEM TEMP'e yazılıp SİLİNMELİ — `--cikti` ile aynı klasöre
    (ör. `_oa/cikti`) bir kopya bırakmak `pipeline_kayit._dilekce_sekilli_
    makbuzsuz_uyarisi`nı YANLIŞLIKLA tetikler (dosya içeriği dilekçe-şekilli
    desenler taşıdığından, taslaktan DAHA YENİ bir 'makbuzsuz aday' gibi
    okunur — canlı saha regresyonu)."""
    girdi = tmp_path / "taslak.md"
    girdi.write_text("Sayın Mahkeme,\n\nNetice-i talep: kabulünü arz ederiz.\n",
                      encoding="utf-8")
    cikti = tmp_path / "taslak.md.udf"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--girdi", str(girdi), "--cikti", str(cikti)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=180)
    assert cp.returncode == 0, cp.stdout + cp.stderr
    kalanlar = sorted(p.name for p in tmp_path.iterdir())
    assert kalanlar == ["taslak.md", "taslak.md.udf"], (
        f"cikti klasöründe beklenmedik artefakt kaldı: {kalanlar}")


@pytest.mark.skipif(not _npx_hazir_mi(), reason="npx/udf-cli oturumu bu makinede kullanılamıyor")
def test_cli_html_acikca_verilirse_silinmez(tmp_path):
    """--html AÇIKÇA verildiğinde bu kullanıcının bilinçli tercihidir —
    dosya KORUNUR (yalnız --html verilmediğinde geçici/silinir)."""
    girdi = tmp_path / "taslak.md"
    girdi.write_text("Metin.\n", encoding="utf-8")
    cikti = tmp_path / "cikti.udf"
    html = tmp_path / "kalici.html"
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--girdi", str(girdi), "--cikti", str(cikti),
         "--html", str(html)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=180)
    assert cp.returncode == 0, cp.stdout + cp.stderr
    assert html.is_file(), "--html AÇIKÇA verildiğinde dosya korunmalıydı"


@pytest.mark.skipif(not _npx_hazir_mi(), reason="npx/udf-cli oturumu bu makinede kullanılamıyor")
def test_cli_kok_ile_goreli_yollari_cozer(tmp_path):
    (tmp_path / "girdi.md").write_text("Metin.\n", encoding="utf-8")
    cp = subprocess.run(
        [sys.executable, str(SCRIPT), "--kok", str(tmp_path),
         "--girdi", "girdi.md", "--cikti", "cikti.udf"],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=180)
    assert cp.returncode == 0, cp.stdout + cp.stderr
    assert (tmp_path / "cikti.udf").is_file()


# ── udf_dogrula(): <tab/> gibi offset TAŞIYAN ama <content> OLMAYAN elemanlar ──

def test_udf_dogrula_tab_elemanli_gercek_belgeyi_GECERLI_sayar():
    """SAHA KANITLI YANLIŞ-BLOK (v0.5.5.2): süreklilik denetimi yalnız
    <content> elemanlarına bakıyordu. Gerçek `udf-cli html2udf` çıktısında
    CDATA'daki bazı karakterler <content> DIŞINDA kendi etiketiyle temsil
    edilir — ölçülen gerçek dilekçede 8 adet `<tab startOffset=".." length="1"/>`.
    Her biri imleci bir karakter kaydırdığından, avukatın UYAP'ta AÇILDIĞINI
    TEYİT ETTİĞİ 46.336 karakterlik dilekçe "offset süreksiz: beklenen 61,
    bulunan 62" ile GEÇERSİZ işaretleniyordu: kapı, korumaya çalıştığı teslimi
    kesiyordu. Denetlenen invaryant 'paragraflar ardışık' değil, 'offset taşıyan
    TÜM elemanlar CDATA'yı boşluksuz/örtüşmesiz döşer'dir."""
    xml_ham = (
        '<?xml version="1.0" encoding="UTF-8" ?>\n'
        '<template format_id="1.8">\n'
        '<content><![CDATA[DOSYA NO\t: 2026/307\n]]></content>'
        '<elements resolver="hvl-default">\n'
        '<paragraph>'
        '<content startOffset="0" length="8"/>'
        '<tab startOffset="8" length="1"/>'
        '<content startOffset="9" length="11"/>'
        '</paragraph>\n'
        '</elements>\n'
        '</template>\n'
    )
    import tempfile
    tmp = tempfile.mktemp(suffix=".udf")
    with zipfile.ZipFile(tmp, "w") as z:
        z.writestr("content.xml", xml_ham.encode("utf-8"))

    sonuc = uy.udf_dogrula(tmp, resmi_okuyucu=False)

    assert sonuc["offsetler_tutarli"] is True, sonuc["hatalar"]
    assert sonuc["gecerli"] is True, sonuc["hatalar"]


def test_udf_dogrula_GERCEK_METIN_kaybini_hala_yakalar():
    """Gevşetme DEĞİL keskinleştirme: kaplanmayan aralıkta METİN varsa hâlâ
    BLOKLAR. Ölçüt "kaplanmayan karakter var mı" değil, "kaplanmayan aralıkta
    METİN var mı"dır — CDATA'da duran ama hiçbir elemanla kaplanmayan metin,
    UYAP'ta görünmeyecek metindir (gerçek kayıp)."""
    xml_ham = (
        '<?xml version="1.0" encoding="UTF-8" ?>\n'
        '<template format_id="1.8">\n'
        '<content><![CDATA[DOSYA NO GIZLI: 2026/307\n]]></content>'
        '<elements resolver="hvl-default">\n'
        '<paragraph>'
        '<content startOffset="0" length="8"/>'
        '<content startOffset="15" length="12"/>'   # 8-15 arası "GIZLI: " KAPLANMIYOR
        '</paragraph>\n'
        '</elements>\n'
        '</template>\n'
    )
    import tempfile
    tmp = tempfile.mktemp(suffix=".udf")
    with zipfile.ZipFile(tmp, "w") as z:
        z.writestr("content.xml", xml_ham.encode("utf-8"))

    sonuc = uy.udf_dogrula(tmp, resmi_okuyucu=False)

    assert sonuc["offsetler_tutarli"] is False
    assert sonuc["gecerli"] is False
    assert any("METİN İÇERİYOR" in h for h in sonuc["hatalar"])


def test_udf_dogrula_editorde_kaydedilmis_bos_paragrafi_GECERLI_sayar():
    """İKİNCİ SAHA VAKASI (v0.5.5.2): UYAP editöründe kaydedilen dosyada BOŞ
    PARAGRAF'ın ayraç newline'ı hiçbir offset'li elemanla kaplanmaz. Avukatın
    imzalayıp UYAP'a yükleyeceği 51.243 karakterlik gerçek dilekçe bu yüzden
    GEÇERSİZ işaretlenmişti — resmî okuyucu ise dosyayı sorunsuz okuyordu.
    Kaplanmayan aralık YALNIZ boşluksa yapısaldır: görünür NOT düşülür,
    teslim KESİLMEZ."""
    xml_ham = (
        '<?xml version="1.0" encoding="UTF-8" ?>\n'
        '<template format_id="1.8">\n'
        '<content><![CDATA[e-imza\n\n  \n]]></content>'
        '<elements resolver="hvl-default">\n'
        '<paragraph><content startOffset="0" length="7"/></paragraph>\n'
        '<paragraph/>\n'                        # boş paragraf: 7. karakter kaplanmaz
        '<paragraph><content startOffset="8" length="3"/></paragraph>\n'
        '</elements>\n'
        '</template>\n'
    )
    import tempfile
    tmp = tempfile.mktemp(suffix=".udf")
    with zipfile.ZipFile(tmp, "w") as z:
        z.writestr("content.xml", xml_ham.encode("utf-8"))

    sonuc = uy.udf_dogrula(tmp, resmi_okuyucu=False)

    assert sonuc["gecerli"] is True, sonuc["hatalar"]
    assert sonuc["kapsanmayan_bosluk"], "boşluk sessizce yutulmamalı, NOT olarak kaydedilmeli"
    assert any("boş paragraf ayracı" in n for n in sonuc["notlar"])


def test_udf_dogrula_imza_dosyasini_bildirir():
    """E-İMZALI nüsha ayırt edilir: arşivde `sign.sgn` varsa bildirilir. Teslim
    öncesi kritiktir — imzalı dosyanın içeriği değişirse imza geçersiz kalır,
    yani o dosya "düzenlenecek taslak" DEĞİL teslim nüshasıdır. Script imzanın
    GEÇERLİLİĞİNİ doğrulamaz (kripto doğrulama UYAP/e-imza altyapısının işidir),
    yalnız VARLIĞINI bildirir — sahte kesinlik yok."""
    xml_ham = (
        '<?xml version="1.0" encoding="UTF-8" ?>\n'
        '<template format_id="1.8">\n'
        '<content><![CDATA[Metin\n]]></content>'
        '<elements resolver="hvl-default">'
        '<paragraph><content startOffset="0" length="6"/></paragraph>'
        '</elements>\n</template>\n'
    )
    import tempfile
    imzali = tempfile.mktemp(suffix=".udf")
    with zipfile.ZipFile(imzali, "w") as z:
        z.writestr("content.xml", xml_ham.encode("utf-8"))
        z.writestr("sign.sgn", b"\x30\x82\x0a\x40sahte-pkcs7-govdesi")
    imzasiz = tempfile.mktemp(suffix=".udf")
    with zipfile.ZipFile(imzasiz, "w") as z:
        z.writestr("content.xml", xml_ham.encode("utf-8"))

    assert uy.udf_dogrula(imzali, resmi_okuyucu=False)["imza_dosyasi"] == "sign.sgn"
    assert uy.udf_dogrula(imzasiz, resmi_okuyucu=False)["imza_dosyasi"] is None
